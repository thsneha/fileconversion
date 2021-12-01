import os
import docx
import glob
from docx.shared import Cm
import re
files = glob.glob(os.getcwd() + "\\james\\translate\\*") # file location
print(files)
for folder in files:
    #print(folder)
    nextpath = glob.glob(folder + "\\*.md")
    print(nextpath)
    doc = docx.Document()
    sections = doc.sections

    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        bookname = folder.split("\\")[-1]
        print(bookname)
        heading = doc.add_heading(bookname,level=1).alignment=1
        print(heading)
        for folderin in sorted(nextpath):
            fin = glob.glob(folderin)
            for fl in sorted(fin):
                title = fl.split("\\")[-3:]
                print(title)
                titlepath="\\".join(title)
                print(titlepath)
                with open(fl,"r",encoding="utf8") as file:
                    content = file.read()
                    print(content)
                    splitcon = content.split("\n")
                    table = doc.add_table(rows = 0,cols =1)
                    table.style = "Table Grid"
                    cells = table.add_row().cells
                    cells[0].paragraphs[0].add_run(titlepath).bold = True

                    table = doc.add_table(rows=0,cols=3)
                    table.style ="Table Grid"
                    table.autofit =True
                    table.allow_autofit = True
                    table.columns[0].width =Cm(2)
                    table.columns[1].width =Cm(9)
                    table.columns[2].width =Cm(9)

                    heading1 = table.add_row().cells
                    heading1[0].paragraphs[0].add_run('NO').bold = True
                    heading1[1].paragraphs[0].add_run('ENGLISH ').bold = True
                    heading1[2].paragraphs[0].add_run('TRANSLATION').bold = True
                    count = 1
                    for lines in splitcon:
                        cells = table.add_row().cells
                        cells[0].paragraphs[0].add_run(str(count)).font.size = Cm(0.5)
                        cells[1].paragraphs[0].add_run(str(lines)).font.size = Cm(0.5)
                        cells[2].text = ''
                        count+=1
                        doc.add_paragraph('')
                    doc.add_page_break()

        doc.save(bookname+ '.docx' )
        print("saved")






